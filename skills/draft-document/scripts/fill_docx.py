#!/usr/bin/env python3
"""雛形(.docx)のプレースホルダーを差し込み値で置換してドラフトを出力するスクリプト。

方針:
- 雛形の体裁（フォント・段落・表）を維持したままプレースホルダーのみを置換する。
  プレースホルダーが複数 run に分割されている場合も、プレースホルダー開始位置の run の
  書式を置換値に引き継ぎ、前後のテキストの書式には手を加えない。
- 本文の段落・表に加え、**入れ子の表（テーブル内テーブル）・ヘッダー・フッター** も
  再帰的に走査する。未置換のプレースホルダーがサイレントに残ることを防ぐため、
  未解決検出も同じ走査ロジックを使う。
- プレースホルダーは ``{{key}}`` 形式（例: ``{{contract_name}}``）。
- 置換値に含まれない項目は ``※要確認`` として残す運用を推奨（本スクリプトは
  未指定キーを既定で ``※要確認：<key>`` に置換する。--keep-unknown で原文維持も可）。

依存: python-docx (``pip install python-docx``)

使い方:
    python fill_docx.py \
        --template assets/amendment-template.docx \
        --output draft.docx \
        --data data.json

data.json 例:
    {
      "contract_name": "業務委託基本契約書",
      "contract_date": "2024年4月1日",
      "party_a": "株式会社サンプル",
      "party_b": "サンプル合同会社",
      "target_article": "第5条（契約期間）",
      "before_text": "本契約の有効期間は...",
      "after_text": "本契約の有効期間は...（1年間延長）",
      "effective_date": "2025年4月1日"
    }

差し込み値は文字列または数値のみを受け付ける。bool・配列・オブジェクトは
契約文書上の表記が定まらないため、エラーとして扱う。
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

try:
    from docx import Document  # type: ignore
except ImportError:  # pragma: no cover
    sys.stderr.write(
        "python-docx が必要です。`pip install python-docx` を実行してください。\n"
    )
    raise

PLACEHOLDER_RE = re.compile(r"\{\{\s*([\w.-]+)\s*\}\}")
# ※要確認のキー名はプレースホルダー由来（ASCII）に限定する。\w を使うと日本語の
# 後続文まで巻き込んでしまうため、明示的な文字クラスにしている。
UNRESOLVED_NOTE_RE = re.compile(r"※要確認：([A-Za-z0-9_.-]+)")


# ---------------------------------------------------------------------------
# 走査（本文・表・入れ子の表・ヘッダー・フッター）
# ---------------------------------------------------------------------------

def _iter_table_paragraphs(table):
    """表内のすべての段落を、入れ子の表も含めて再帰的に走査する。"""
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def _iter_paragraphs(document):
    """文書内のすべての段落を走査する（本文・表・ヘッダー・フッター）。"""
    yield from document.paragraphs
    for table in document.tables:
        yield from _iter_table_paragraphs(table)
    for section in document.sections:
        for part in (section.header, section.footer):
            yield from part.paragraphs
            for table in part.tables:
                yield from _iter_table_paragraphs(table)


# ---------------------------------------------------------------------------
# 置換（run の書式を維持）
# ---------------------------------------------------------------------------

def _substitution(match: re.Match, data: dict, keep_unknown: bool) -> str | None:
    """プレースホルダーの置換文字列を返す。None は「置換しない（原文維持）」。"""
    key = match.group(1)
    value = data.get(key)
    if value is not None and value != "":
        return str(value)
    return None if keep_unknown else f"※要確認：{key}"


def _splice_runs(runs, start: int, end: int, replacement: str) -> None:
    """段落内 run 群の文字位置 [start, end) を replacement で置き換える。

    置換値はプレースホルダー開始位置の run に書き込むため、その run の書式
    （太字・斜体等）が置換値に引き継がれる。範囲外のテキストの書式は変更しない。
    """
    pos = 0
    spans = []
    for run in runs:
        spans.append((pos, pos + len(run.text), run))
        pos += len(run.text)

    inserted = False
    for span_start, span_end, run in spans:
        if span_end <= start or span_start >= end:
            continue  # 置換範囲と重ならない run はそのまま
        head = run.text[: max(start - span_start, 0)]
        tail = run.text[end - span_start:] if end - span_start < len(run.text) else ""
        if not inserted:
            run.text = head + replacement + tail
            inserted = True
        else:
            run.text = head + tail


def _replace_in_paragraph(paragraph, data: dict, keep_unknown: bool) -> None:
    """段落内のプレースホルダーを、run の書式を維持したまま置換する。"""
    full_text = "".join(run.text for run in paragraph.runs)
    if "{{" not in full_text:
        return
    # 後ろのマッチから処理することで、置換による文字位置のずれを回避する
    for match in reversed(list(PLACEHOLDER_RE.finditer(full_text))):
        replacement = _substitution(match, data, keep_unknown)
        if replacement is None:
            continue
        _splice_runs(paragraph.runs, match.start(), match.end(), replacement)


# ---------------------------------------------------------------------------
# 入出力
# ---------------------------------------------------------------------------

def validate_data(data) -> None:
    """差し込みデータを検証する。不正な型はエラーにする（無条件の str() 変換をしない）。"""
    if not isinstance(data, dict):
        raise ValueError(
            "差し込みデータ(JSON)は {\"key\": \"値\"} 形式のオブジェクトである必要があります。"
        )
    invalid = [
        key
        for key, value in data.items()
        if isinstance(value, bool)
        or not isinstance(value, (str, int, float, type(None)))
    ]
    if invalid:
        raise ValueError(
            "差し込み値は文字列または数値のみ使用できます。"
            f"次のキーの値を修正してください: {', '.join(sorted(invalid))}"
        )


def fill_template(template: Path, output: Path, data: dict, keep_unknown: bool) -> list[str]:
    """雛形を差し込み、未解決プレースホルダーの一覧を返す。"""
    try:
        document = Document(str(template))
    except Exception as exc:  # noqa: BLE001 - 利用者向けメッセージに変換
        raise ValueError(f"雛形を開けませんでした（{template}）: {exc}") from exc

    for paragraph in _iter_paragraphs(document):
        _replace_in_paragraph(paragraph, data, keep_unknown)

    # 未解決（{{...}} または ※要確認）の検出。置換と同じ走査ロジックを使うことで、
    # 入れ子の表・ヘッダー・フッター内の未置換もサイレントに残さない。
    unresolved: list[str] = []
    for paragraph in _iter_paragraphs(document):
        text = paragraph.text
        for key in PLACEHOLDER_RE.findall(text):
            unresolved.append(f"{{{{{key}}}}}")
        for key in UNRESOLVED_NOTE_RE.findall(text):
            unresolved.append(f"※要確認：{key}")

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))
    return sorted(set(unresolved))


def main() -> int:
    parser = argparse.ArgumentParser(description="法務文書ドラフトの雛形差し込み")
    parser.add_argument("--template", required=True, type=Path, help="雛形 .docx のパス")
    parser.add_argument("--output", required=True, type=Path, help="出力 .docx のパス")
    parser.add_argument("--data", required=True, type=Path, help="差し込み値の JSON パス")
    parser.add_argument(
        "--keep-unknown",
        action="store_true",
        help="未指定キーを ※要確認 に変換せず {{key}} のまま残す",
    )
    args = parser.parse_args()

    if not args.template.is_file():
        sys.stderr.write(f"雛形ファイルが見つかりません: {args.template}\n")
        return 1
    if not args.data.is_file():
        sys.stderr.write(f"差し込みデータ(JSON)が見つかりません: {args.data}\n")
        return 1

    try:
        data = json.loads(args.data.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        sys.stderr.write(f"差し込みデータ(JSON)を読み込めませんでした: {exc}\n")
        return 1

    try:
        validate_data(data)
        unresolved = fill_template(args.template, args.output, data, args.keep_unknown)
    except ValueError as exc:
        sys.stderr.write(f"{exc}\n")
        return 1

    print(f"ドラフトを出力しました: {args.output}")
    if unresolved:
        print("\n【確認事項（未解決プレースホルダー）】")
        for item in unresolved:
            print(f"  - {item}")
        print("\n上記は推測で埋めていません。法務担当者による確認・入力が必要です。")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
