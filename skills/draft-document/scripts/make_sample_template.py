#!/usr/bin/env python3
"""汎用ダミーの変更覚書雛形 (.docx) を生成するスクリプト。

`assets/amendment-template.docx` を再生成したいときに使う。プレースホルダーは
``{{key}}`` 形式で、`fill_docx.py` で差し込める。同梱する雛形はすべて汎用ダミーであり、
特定顧客・組織の情報は含まない。

依存: python-docx (``pip install python-docx``)

使い方:
    python make_sample_template.py --output assets/amendment-template.docx
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

try:
    from docx import Document  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
except ImportError:  # pragma: no cover
    sys.stderr.write(
        "python-docx が必要です。`pip install python-docx` を実行してください。\n"
    )
    raise


def build_template(output: Path) -> None:
    doc = Document()

    note = doc.add_paragraph()
    run = note.add_run("※ 本文書はドラフト（草案）です。確定前に法務担当者のレビューを受けてください。")
    run.italic = True

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("{{contract_name}} 変更覚書").bold = True

    doc.add_paragraph(
        "{{party_a}}（以下「甲」という。）と{{party_b}}（以下「乙」という。）とは、"
        "甲乙間で {{contract_date}} に締結した「{{contract_name}}」（以下「原契約」という。）"
        "について、以下のとおり変更することに合意し、本変更覚書（以下「本覚書」という。）を締結する。"
    )

    doc.add_paragraph("第1条（原契約の変更）")
    doc.add_paragraph("  原契約{{target_article}}を次のとおり変更する。")
    doc.add_paragraph("  【変更前】")
    doc.add_paragraph("  {{before_text}}")
    doc.add_paragraph("  【変更後】")
    doc.add_paragraph("  {{after_text}}")

    doc.add_paragraph("第2条（効力発生日）")
    doc.add_paragraph("  本覚書は、{{effective_date}}から効力を生じる。")

    doc.add_paragraph("第3条（原契約の効力）")
    doc.add_paragraph("  本覚書に定めのない事項については、原契約の定めをそのまま適用する。")

    doc.add_paragraph(
        "本覚書の成立を証するため、本書2通を作成し、甲乙記名押印のうえ、各1通を保有する。"
    )
    doc.add_paragraph("{{execution_date}}")
    doc.add_paragraph("甲：{{party_a}}　　　　　　　　　　印")
    doc.add_paragraph("乙：{{party_b}}　　　　　　　　　　印")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    print(f"雛形を生成しました: {output}")


def main() -> int:
    parser = argparse.ArgumentParser(description="汎用ダミー変更覚書雛形の生成")
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("assets/amendment-template.docx"),
        help="出力先 .docx パス",
    )
    args = parser.parse_args()
    build_template(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
