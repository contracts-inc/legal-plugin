#!/usr/bin/env python3
"""fill_docx.py のユニットテスト。

実行方法（scripts/ ディレクトリで）:
    python -m unittest test_fill_docx -v

カバーするケース:
- 単一 run の置換
- 複数 run に分割されたプレースホルダーの置換（書式維持）
- 表内の置換
- 入れ子の表（テーブル内テーブル）内の置換と未解決検出
- 未指定キーの ※要確認 変換 / --keep-unknown 相当の原文維持
- 差し込みデータの型検証
"""
from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

import fill_docx


class FillDocxTest(unittest.TestCase):
    def _run_fill(self, build, data, keep_unknown=False):
        """build(doc) で雛形を組み立て、差し込み結果の Document と未解決一覧を返す。"""
        with tempfile.TemporaryDirectory() as tmp:
            template = Path(tmp) / "template.docx"
            output = Path(tmp) / "output.docx"
            doc = Document()
            build(doc)
            doc.save(str(template))
            unresolved = fill_docx.fill_template(template, output, data, keep_unknown)
            return Document(str(output)), unresolved

    # ------------------------------------------------------------------
    # 置換の基本動作
    # ------------------------------------------------------------------

    def test_single_run(self):
        doc, unresolved = self._run_fill(
            lambda d: d.add_paragraph("契約名: {{contract_name}}"),
            {"contract_name": "業務委託基本契約書"},
        )
        self.assertEqual(doc.paragraphs[0].text, "契約名: 業務委託基本契約書")
        self.assertEqual(unresolved, [])

    def test_multi_run_placeholder_keeps_formatting(self):
        """プレースホルダーが複数 run に分割されていても置換され、書式が維持される。"""

        def build(d):
            p = d.add_paragraph()
            p.add_run("契約名: ")
            bold_run = p.add_run("{{contract_")
            bold_run.bold = True
            bold_run.italic = True
            tail_run = p.add_run("name}}")
            tail_run.bold = True
            tail_run.italic = True

        doc, unresolved = self._run_fill(build, {"contract_name": "業務委託基本契約書"})
        p = doc.paragraphs[0]
        self.assertEqual(p.text, "契約名: 業務委託基本契約書")
        self.assertEqual(unresolved, [])
        # 置換値はプレースホルダー開始 run の書式（太字・斜体）を引き継ぐ
        value_runs = [r for r in p.runs if "業務委託基本契約書" in r.text]
        self.assertEqual(len(value_runs), 1)
        self.assertTrue(value_runs[0].bold)
        self.assertTrue(value_runs[0].italic)
        # 前置きテキストの書式は変わらない
        self.assertIsNone(p.runs[0].bold)

    def test_surrounding_text_preserved(self):
        """同一 run 内のプレースホルダー前後のテキストが保持される。"""
        doc, _ = self._run_fill(
            lambda d: d.add_paragraph("甲は{{party_a}}とする。"),
            {"party_a": "株式会社サンプル"},
        )
        self.assertEqual(doc.paragraphs[0].text, "甲は株式会社サンプルとする。")

    # ------------------------------------------------------------------
    # 表・入れ子の表
    # ------------------------------------------------------------------

    def test_table(self):
        def build(d):
            table = d.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "当事者: {{party_a}}"

        doc, unresolved = self._run_fill(build, {"party_a": "株式会社サンプル"})
        self.assertEqual(doc.tables[0].cell(0, 0).text, "当事者: 株式会社サンプル")
        self.assertEqual(unresolved, [])

    def test_nested_table_replaced_and_detected(self):
        """入れ子の表も置換対象になり、未指定キーは ※要確認 として検出される。"""

        def build(d):
            outer = d.add_table(rows=1, cols=1)
            outer.cell(0, 0).text = "外側セル: {{outer_key}}"
            inner = outer.cell(0, 0).add_table(rows=1, cols=1)
            inner.cell(0, 0).text = "内側セル: {{inner_key}}"

        doc, unresolved = self._run_fill(build, {"outer_key": "外側の値"})
        outer_cell = doc.tables[0].cell(0, 0)
        inner_cell = outer_cell.tables[0].cell(0, 0)
        self.assertIn("外側の値", outer_cell.paragraphs[0].text)
        self.assertEqual(inner_cell.text, "内側セル: ※要確認：inner_key")
        self.assertIn("※要確認：inner_key", unresolved)

    # ------------------------------------------------------------------
    # 未解決の扱い
    # ------------------------------------------------------------------

    def test_unknown_key_becomes_note(self):
        doc, unresolved = self._run_fill(
            lambda d: d.add_paragraph("効力発生日: {{effective_date}}"), {}
        )
        self.assertEqual(doc.paragraphs[0].text, "効力発生日: ※要確認：effective_date")
        self.assertEqual(unresolved, ["※要確認：effective_date"])

    def test_keep_unknown(self):
        doc, unresolved = self._run_fill(
            lambda d: d.add_paragraph("効力発生日: {{effective_date}}"),
            {},
            keep_unknown=True,
        )
        self.assertEqual(doc.paragraphs[0].text, "効力発生日: {{effective_date}}")
        self.assertEqual(unresolved, ["{{effective_date}}"])

    # ------------------------------------------------------------------
    # 入力検証
    # ------------------------------------------------------------------

    def test_validate_data_rejects_bool_and_containers(self):
        for bad in ({"flag": True}, {"items": ["a"]}, {"nested": {"k": "v"}}):
            with self.assertRaises(ValueError):
                fill_docx.validate_data(bad)

    def test_validate_data_accepts_str_and_numbers(self):
        fill_docx.validate_data({"name": "値", "count": 1, "rate": 0.5, "skip": None})

    def test_missing_template_raises_value_error(self):
        with tempfile.TemporaryDirectory() as tmp:
            with self.assertRaises(ValueError):
                fill_docx.fill_template(
                    Path(tmp) / "not-exist.docx", Path(tmp) / "out.docx", {}, False
                )


if __name__ == "__main__":
    unittest.main()
